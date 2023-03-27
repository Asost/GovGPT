from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import time

class GenerateDocx:
	def replaceTerm(self, text, term):
		term_len = len(term)
		term_index = text.lower().find(term.lower())
		if term_index != -1:
			text = text[:term_index] + "_" * term_len + text[term_index + term_len:]
		return text

	def __init__(self, sgInfo, chapterNumber, formatedResponse, questionless=True, name="John Doe", date=date.today().strftime('%B %d, %Y'), period="6"):
		self.sgInfo = sgInfo
		self.chapterNumber = chapterNumber
		self.formatedResponse = formatedResponse
		self.questionless = questionless
		self.name = name
		self.date = date
		self.period = period

	def generateQuestionless(self):
		# Create final document
		finalDoc = Document()

		# Add the header (required for assignmnt points)
		header = finalDoc.sections[0].header
		paragraph = header.add_paragraph('Name: ' + self.name + '\nDate: ' + self.date + '\n Period: ' + str(self.period))
		paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

		# Parse and add questions to document
		#questionList = []
		for i, answer in enumerate(self.formatedResponse["questions"]):
			question = self.sgInfo["questions"][i]
			response = self.formatedResponse["synonymAnswers"][i]["response"]
			related = self.formatedResponse["synonymAnswers"][i]["related"]
			# strip together answer and add to document
			finalResponse = finalDoc.add_paragraph("", style='List Number')
			finalResponse.add_run('\n\t' + response.lstrip('\n') + ' ')
			run = finalResponse.add_run(related.lstrip('\n'))
			run.bold = True
			run.underline = True

			# have each question paragraph object be appended to the list
			#questionList.append(finalResponse)

		keyTermsAnnouncement = finalDoc.add_paragraph()
		run = keyTermsAnnouncement.add_run('KEY TERMS: ')
		run.bold = True

		# Parse and add Key Terms to document
		#termsList = []
		
		for i, answer in enumerate(self.formatedResponse["terms"]):
			response = self.replaceTerm(answer["response"], answer["term"])
			# strip together answer and add to document
			finalResponse = finalDoc.add_paragraph(style='List Number 2')
			run = finalResponse.add_run()
			run.underline = True
			finalResponse.add_run(response.lstrip('\n'))

			# have each term paragraph object be appended to the list
			#termsList.append(finalResponse)
		return finalDoc

	# Create a Document object for the final results
	def generateAIdocx(self):
		# Create final document
		finalDoc = Document()

		# Add the header (required for assignmnt points)
		header = finalDoc.sections[0].header
		paragraph = header.add_paragraph('Name: ' + self.name + '\nDate: ' + self.date + '\n Period: ' + str(self.period))
		paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

		# Add the title
		title = finalDoc.add_paragraph()
		title.alignment = WD_ALIGN_PARAGRAPH.CENTER
		run = title.add_run('A.P. GOVERNMENT & POLITICS\nCHAPTER ' + str(int(self.chapterNumber)) + '\n' + self.sgInfo["title"])
		run.bold = True

		# Parse and add questions to document
		#questionList = []
		for i, answer in enumerate(self.formatedResponse["questions"]):
			question = self.sgInfo["questions"][i]
			response = self.formatedResponse["synonymAnswers"][i]["response"]
			related = self.formatedResponse["synonymAnswers"][i]["related"]
			# strip together answer and add to document
			finalResponse = finalDoc.add_paragraph(question, style='List Number')
			finalResponse.add_run('\n\t' + response.lstrip('\n') + ' ')
			run = finalResponse.add_run(related.lstrip('\n'))
			run.bold = True
			run.underline = True

			# have each question paragraph object be appended to the list
			#questionList.append(finalResponse)

		keyTermsAnnouncement = finalDoc.add_paragraph()
		run = keyTermsAnnouncement.add_run('KEY TERMS: ')
		run.bold = True
		keyTermsAnnouncement.add_run('(Define in complete Sentences)')

		# Parse and add Key Terms to document
		#termsList = []
		
		for i, answer in enumerate(self.formatedResponse["terms"]):
			# strip together answer and add to document
			finalResponse = finalDoc.add_paragraph(style='List Number 2')
			run = finalResponse.add_run(answer["term"])
			run.underline = True
			finalResponse.add_run(": " + answer["response"].lstrip('\n'))

			# have each term paragraph object be appended to the list
			#termsList.append(finalResponse)
		
		if self.questionless:
			questionlessDocument = self.generateQuestionless()
			return finalDoc, questionlessDocument
		return finalDoc, None