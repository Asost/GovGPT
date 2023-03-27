import requests
import os
from io import BytesIO
from docx import Document

class Parser:
	# Get the url given and turn into a Document object
	def getFileAsDoc(self):
		if os.path.isfile(self.url):
			return Document(self.url)
		else:
			return Document(BytesIO(requests.get(self.url).content))

	def __init__(self, url):
		self.url = url
		self.doc = self.getFileAsDoc()

	# Parse the document for the information
	def parse(self):
		sgInfo = {"title": "", "questions": [], "keyTerms": []}
		keyTerms = False
		for i, paragraph in enumerate(self.doc.paragraphs):
			if(i == 2):
				sgInfo["title"] = paragraph.text
			elif("KEY TERMS" in paragraph.text):
				keyTerms = True
			elif(i >= 4 and not keyTerms):
				sgInfo["questions"].append(paragraph.text)
			elif(keyTerms):
				sgInfo["keyTerms"].append(paragraph.text)
		sgInfo["questions"] = list(filter(None, sgInfo["questions"]))
		sgInfo["keyTerms"] = list(filter(None, sgInfo["keyTerms"]))

		self.sgInfo = sgInfo
		return sgInfo